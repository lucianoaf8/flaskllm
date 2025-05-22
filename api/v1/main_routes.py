# api/v1/__init__.py
from flask import Blueprint
from flask_restx import Api

# Create blueprint
v1_blueprint = Blueprint('v1', __name__, url_prefix='/api/v1')

# Create API with Swagger documentation
api = Api(
    v1_blueprint,
    version='1.0',
    title='FlaskLLM API',
    description='A secure, scalable API for processing prompts with LLMs',
    doc='/docs',
    default='LLM',
    default_label='LLM API Operations',
    validate=True,
)

# llm/anthropic_handler.py
from typing import Generator, Optional, Union

def process_prompt(
    self,
    prompt: str,
    source: Optional[str] = None,
    language: Optional[str] = None,
    type: Optional[str] = None,
    stream: bool = False,
    **kwargs
) -> Union[str, Generator[str, None, None]]:
    """
    Process a prompt using the Anthropic API.

    Args:
        prompt: The prompt to process
        source: Source of the prompt (email, meeting, etc.)
        language: Target language code (ISO 639-1)
        type: Type of processing to perform
        stream: Whether to stream the response
        **kwargs: Additional parameters for the Anthropic API

    Returns:
        Processed result as a string or a generator yielding chunks of the response

    Raises:
        LLMAPIError: If the Anthropic API returns an error
    """
    try:
        # Create system prompt
        system_prompt = self._create_system_prompt(source, language, type)

        # Log the request (without the full prompt for privacy)
        logger.info(
            "Sending request to Anthropic",
            model=self.model,
            prompt_length=len(prompt),
            source=source,
            language=language,
            type=type,
            stream=stream
        )

        # Prepare parameters
        params = {
            "model": self.model,
            "system": system_prompt,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": kwargs.get("max_tokens", self.max_tokens_to_sample),
            "temperature": kwargs.get("temperature", 0.3),
            "stream": stream
        }

        # Add other supported parameters
        for param in ["top_k", "top_p", "stop_sequences"]:
            if param in kwargs:
                params[param] = kwargs[param]

        # Send request to Anthropic
        response = self.client.messages.create(**params)

        if stream:
            # Return a generator for streaming responses
            def response_generator():
                collected_chunks = []
                # Process the streaming response
                for chunk in response:
                    if not chunk.delta or not chunk.delta.text:
                        continue

                    content = chunk.delta.text
                    if content:
                        collected_chunks.append(content)
                        yield content

                # Log the complete response after streaming
                logger.info(
                    "Completed streaming response",
                    total_chunks=len(collected_chunks),
                    total_length=sum(len(c) for c in collected_chunks)
                )

            return response_generator()
        else:
            # Extract and return the response content for non-streaming
            if not response.content:
                raise LLMAPIError("Empty response from Anthropic API")

            # Get the first content block
            content_block = response.content[0]

            # Check if it's a text block and has text content
            if content_block.type != "text" or not content_block.text:
                raise LLMAPIError("Invalid or empty response from Anthropic API")

            # Text is now guaranteed to be a string
            result = cast(str, content_block.text)
            return result

    except anthropic.AuthenticationError as e:
        logger.error("Anthropic authentication error", error=str(e))
        raise LLMAPIError("Authentication error with Anthropic API")

    except anthropic.RateLimitError as e:
        logger.error("Anthropic rate limit exceeded", error=str(e))
        raise LLMAPIError("Rate limit exceeded with Anthropic API")

    except anthropic.APIError as e:
        logger.error("Anthropic API error", error=str(e))
        raise LLMAPIError(f"Error from Anthropic API: {str(e)}")

    except Exception as e:
        logger.exception("Unexpected error with Anthropic API", error=str(e))
        raise LLMAPIError(f"Unexpected error: {str(e)}")

# llm/template_storage.py
import json
import os
from typing import Dict, List, Optional

from core.exceptions import TemplateError
from core.logging import get_logger
from llm.templates import PromptTemplate

logger = get_logger(__name__)

class TemplateStorage:
    """
    Storage for prompt templates.

    Supports loading templates from files and in-memory storage.
    """

    def __init__(self, templates_dir: Optional[str] = None):
        """
        Initialize the template storage.

        Args:
            templates_dir: Directory containing template JSON files
        """
        self.templates: Dict[str, PromptTemplate] = {}

        # Load templates from directory if provided
        if templates_dir and os.path.isdir(templates_dir):
            self._load_templates_from_dir(templates_dir)

    def _load_templates_from_dir(self, templates_dir: str) -> None:
        """
        Load templates from JSON files in a directory.

        Args:
            templates_dir: Directory containing template JSON files
        """
        for filename in os.listdir(templates_dir):
            if filename.endswith('.json'):
                file_path = os.path.join(templates_dir, filename)
                try:
                    with open(file_path, 'r') as f:
                        template_data = json.load(f)

                    template = PromptTemplate(**template_data)
                    self.templates[template.id] = template
                    logger.info(f"Loaded template {template.id} from {filename}")
                except Exception as e:
                    logger.error(f"Error loading template from {filename}: {str(e)}")

    def get_template(self, template_id: str) -> PromptTemplate:
        """
        Get a template by ID.

        Args:
            template_id: Template ID

        Returns:
            Prompt template

        Raises:
            TemplateError: If the template is not found
        """
        if template_id not in self.templates:
            raise TemplateError(f"Template not found: {template_id}")

        return self.templates[template_id]

    def add_template(self, template: PromptTemplate) -> None:
        """
        Add a template to the storage.

        Args:
            template: Prompt template

        Raises:
            TemplateError: If a template with the same ID already exists
        """
        if template.id in self.templates:
            raise TemplateError(f"Template with ID {template.id} already exists")

        self.templates[template.id] = template
        logger.info(f"Added template {template.id}")

    def update_template(self, template: PromptTemplate) -> None:
        """
        Update a template in the storage.

        Args:
            template: Prompt template

        Raises:
            TemplateError: If the template is not found
        """
        if template.id not in self.templates:
            raise TemplateError(f"Template not found: {template.id}")

        self.templates[template.id] = template
        logger.info(f"Updated template {template.id}")

    def delete_template(self, template_id: str) -> None:
        """
        Delete a template from the storage.

        Args:
            template_id: Template ID

        Raises:
            TemplateError: If the template is not found
        """
        if template_id not in self.templates:
            raise TemplateError(f"Template not found: {template_id}")

        del self.templates[template_id]
        logger.info(f"Deleted template {template_id}")

    def list_templates(self) -> List[Dict]:
        """
        List all templates.

        Returns:
            List of template summaries
        """
        return [
            {
                'id': t.id,
                'name': t.name,
                'description': t.description,
                'type': t.type,
                'source': t.source
            }
            for t in self.templates.values()
        ]

# utils/file_processor.py
"""
File Processing Module

This module provides functionality for processing different file types.
"""
import base64
import io
import os
from enum import Enum
from typing import Dict, List, Optional, Union, BinaryIO, Tuple

import docx
import fitz  # PyMuPDF
import pandas as pd
from PIL import Image
import pytesseract
from pydantic import BaseModel

from core.exceptions import FileProcessingError
from core.logging import get_logger

logger = get_logger(__name__)

class FileType(str, Enum):
    """File types supported for processing."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    CSV = "csv"
    XLSX = "xlsx"
    JPG = "jpg"
    PNG = "png"
    UNKNOWN = "unknown"

class FileMetadata(BaseModel):
    """Metadata about a processed file."""
    filename: str
    file_type: FileType
    size_bytes: int
    num_pages: Optional[int] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None

class FileContents(BaseModel):
    """Contents extracted from a file."""
    text: str
    metadata: FileMetadata
    tables: Optional[List[Dict]] = None
    images: Optional[List[str]] = None

class FileProcessor:
    """
    Processor for different file types.

    Handles extraction of text, metadata, tables, and images from various file formats.
    """

    def __init__(self, max_file_size_mb: int = 10):
        """
        Initialize the file processor.

        Args:
            max_file_size_mb: Maximum file size in MB
        """
        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024

    def process_file(self, file: Union[BinaryIO, bytes, str], filename: str) -> FileContents:
        """
        Process a file and extract its contents.

        Args:
            file: File object, file path, or file content as bytes
            filename: Original filename

        Returns:
            Extracted file contents

        Raises:
            FileProcessingError: If file processing fails
        """
        try:
            # Determine file type
            file_type = self._get_file_type(filename)

            # Get file content as bytes
            if isinstance(file, str):
                # File path
                if not os.path.exists(file):
                    raise FileProcessingError(f"File not found: {file}")

                with open(file, 'rb') as f:
                    file_content = f.read()
            elif hasattr(file, 'read'):
                # File-like object
                file_content = file.read()
                if hasattr(file, 'seek'):
                    file.seek(0)  # Reset file pointer for potential reuse
            else:
                # Bytes
                file_content = file

            # Check file size
            if len(file_content) > self.max_file_size_bytes:
                raise FileProcessingError(
                    f"File size ({len(file_content) / 1024 / 1024:.2f} MB) exceeds maximum allowed size "
                    f"({self.max_file_size_bytes / 1024 / 1024:.2f} MB)"
                )

            # Process file based on type
            if file_type == FileType.PDF:
                return self._process_pdf(file_content, filename)
            elif file_type == FileType.DOCX:
                return self._process_docx(file_content, filename)
            elif file_type == FileType.TXT:
                return self._process_txt(file_content, filename)
            elif file_type == FileType.CSV:
                return self._process_csv(file_content, filename)
            elif file_type == FileType.XLSX:
                return self._process_xlsx(file_content, filename)
            elif file_type in [FileType.JPG, FileType.PNG]:
                return self._process_image(file_content, filename, file_type)
            else:
                raise FileProcessingError(f"Unsupported file type: {file_type}")

        except Exception as e:
            logger.error(f"Error processing file {filename}: {str(e)}", exc_info=True)
            raise FileProcessingError(f"Failed to process file: {str(e)}")

    def _get_file_type(self, filename: str) -> FileType:
        """
        Determine file type from filename.

        Args:
            filename: Original filename

        Returns:
            File type
        """
        extension = filename.lower().split('.')[-1] if '.' in filename else ''

        if extension == 'pdf':
            return FileType.PDF
        elif extension == 'docx':
            return FileType.DOCX
        elif extension == 'txt':
            return FileType.TXT
        elif extension == 'csv':
            return FileType.CSV
        elif extension in ['xlsx', 'xls']:
            return FileType.XLSX
        elif extension in ['jpg', 'jpeg']:
            return FileType.JPG
        elif extension == 'png':
            return FileType.PNG
        else:
            return FileType.UNKNOWN

    def _process_pdf(self, content: bytes, filename: str) -> FileContents:
        """
        Process a PDF file.

        Args:
            content: PDF file content
            filename: Original filename

        Returns:
            Extracted file contents
        """
        # Create a memory buffer
        pdf_buffer = io.BytesIO(content)

        # Open the PDF
        try:
            doc = fitz.open(stream=pdf_buffer, filetype="pdf")

            # Extract text from all pages
            text = ""
            images = []

            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()



## Response:
```python
                # Extract images if available
                for img_index, img in enumerate(page.get_images(full=True)):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]

                    # Convert to base64 for storage
                    b64_image = base64.b64encode(image_bytes).decode()
                    images.append(b64_image)

            # Count words and characters
            word_count = len(text.split())
            character_count = len(text)

            # Create metadata
            metadata = FileMetadata(
                filename=filename,
                file_type=FileType.PDF,
                size_bytes=len(content),
                num_pages=len(doc),
                page_count=len(doc),
                word_count=word_count,
                character_count=character_count
            )

            # Create result
            return FileContents(
                text=text,
                metadata=metadata,
                images=images if images else None
            )

        finally:
            # Close the PDF
            if 'doc' in locals():
                doc.close()

    def _process_docx(self, content: bytes, filename: str) -> FileContents:
        """
        Process a DOCX file.

        Args:
            content: DOCX file content
            filename: Original filename

        Returns:
            Extracted file contents
        """
        # Create a memory buffer
        docx_buffer = io.BytesIO(content)

        # Open the DOCX
        doc = docx.Document(docx_buffer)

        # Extract text
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        # Count words and characters
        word_count = len(text.split())
        character_count = len(text)

        # Create metadata
        metadata = FileMetadata(
            filename=filename,
            file_type=FileType.DOCX,
            size_bytes=len(content),
            page_count=len(doc.sections),
            word_count=word_count,
            character_count=character_count
        )

        # Create result
        return FileContents(
            text=text,
            metadata=metadata,
            tables=tables if tables else None
        )

    def _process_txt(self, content: bytes, filename: str) -> FileContents:
        """
        Process a text file.

        Args:
            content: Text file content
            filename: Original filename

        Returns:
            Extracted file contents
        """
        # Decode text
        try:
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            # Try different encoding
            try:
                text = content.decode('latin-1')
            except:
                raise FileProcessingError("Unable to decode text file")

        # Count words and characters
        word_count = len(text.split())
        character_count = len(text)

        # Count lines
        lines = text.split('\n')
        line_count = len(lines)

        # Create metadata
        metadata = FileMetadata(
            filename=filename,
            file_type=FileType.TXT,
            size_bytes=len(content),
            page_count=line_count // 50,  # Rough estimation
            word_count=word_count,
            character_count=character_count
        )

        # Create result
        return FileContents(
            text=text,
            metadata=metadata
        )

    def _process_csv(self, content: bytes, filename: str) -> FileContents:
        """
        Process a CSV file.

        Args:
            content: CSV file content
            filename: Original filename

        Returns:
            Extracted file contents
        """
        # Create a memory buffer
        csv_buffer = io.BytesIO(content)

        # Read CSV
        try:
            df = pd.read_csv(csv_buffer)

            # Convert to text representation
            text = df.to_string()

            # Get table data
            table_data = df.to_dict(orient='records')

            # Create metadata
            metadata = FileMetadata(
                filename=filename,
                file_type=FileType.CSV,
                size_bytes=len(content),
                word_count=len(text.split()),
                character_count=len(text)
            )

            # Create result
            return FileContents(
                text=text,
                metadata=metadata,
                tables=[table_data]
            )

        except Exception as e:
            raise FileProcessingError(f"Failed to parse CSV file: {str(e)}")

    def _process_xlsx(self, content: bytes, filename: str) -> FileContents:
        """
        Process an Excel file.

        Args:
            content: Excel file content
            filename: Original filename

        Returns:
            Extracted file contents
        """
        # Create a memory buffer
        xlsx_buffer = io.BytesIO(content)

        # Read Excel
        try:
            xlsx = pd.ExcelFile(xlsx_buffer)

            all_text = ""
            all_tables = []

            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)

                # Add sheet name and data to text
                sheet_text = f"--- Sheet: {sheet_name} ---\n"
                sheet_text += df.to_string()
                all_text += sheet_text + "\n\n"

                # Add table data
                table_data = {
                    'sheet_name': sheet_name,
                    'data': df.to_dict(orient='records')
                }
                all_tables.append(table_data)

            # Create metadata
            metadata = FileMetadata(
                filename=filename,
                file_type=FileType.XLSX,
                size_bytes=len(content),
                num_pages=len(xlsx.sheet_names),
                word_count=len(all_text.split()),
                character_count=len(all_text)
            )

            # Create result
            return FileContents(
                text=all_text,
                metadata=metadata,
                tables=all_tables
            )

        except Exception as e:
            raise FileProcessingError(f"Failed to parse Excel file: {str(e)}")

    def _process_image(self, content: bytes, filename: str, file_type: FileType) -> FileContents:
        """
        Process an image file using OCR.

        Args:
            content: Image file content
            filename: Original filename
            file_type: Image file type

        Returns:
            Extracted file contents
        """
        # Create a memory buffer
        img_buffer = io.BytesIO(content)

        # Open the image
        try:
            img = Image.open(img_buffer)

            # Perform OCR
            text = pytesseract.image_to_string(img)

            # Convert image to base64 for storage
            img_buffer.seek(0)
            img_base64 = base64.b64encode(img_buffer.read()).decode()

            # Create metadata
            metadata = FileMetadata(
                filename=filename,
                file_type=file_type,
                size_bytes=len(content),
                word_count=len(text.split()),
                character_count=len(text)
            )

            # Create result
            return FileContents(
                text=text,
                metadata=metadata,
                images=[img_base64]
            )

        except Exception as e:
            raise FileProcessingError(f"Failed to process image: {str(e)}")

# llm/conversation_storage.py
import json
import os
from datetime import datetime
from typing import Dict, List, Optional

from core.exceptions import ConversationError
from core.logging import get_logger
from llm.conversation import Conversation, Message, MessageRole

logger = get_logger(__name__)

class ConversationStorage:
    """
    Storage for conversations.

    Supports in-memory storage and file-based persistence.
    """

    def __init__(self, storage_dir: Optional[str] = None):
        """
        Initialize the conversation storage.

        Args:
            storage_dir: Directory for file-based storage
        """
        self.conversations: Dict[str, Conversation] = {}
        self.storage_dir = storage_dir

        # Create storage directory if it doesn't exist
        if storage_dir and not os.path.exists(storage_dir):
            os.makedirs(storage_dir)
            logger.info(f"Created conversation storage directory: {storage_dir}")

    def create_conversation(
        self,
        user_id: str,
        title: Optional[str] = None,
        system_prompt: Optional[str] = None,
        metadata: Optional[Dict] = None
    ) -> Conversation:
        """
        Create a new conversation.

        Args:
            user_id: User ID
            title: Conversation title
            system_prompt: System prompt for the conversation
            metadata: Additional metadata

        Returns:
            Created conversation
        """
        conversation = Conversation(
            user_id=user_id,
            title=title,
            system_prompt=system_prompt,
            metadata=metadata or {}
        )

        self.conversations[conversation.id] = conversation
        self._save_conversation(conversation)

        logger.info(f"Created conversation {conversation.id} for user {user_id}")
        return conversation

    def get_conversation(self, conversation_id: str) -> Conversation:
        """
        Get a conversation by ID.

        Args:
            conversation_id: Conversation ID

        Returns:
            Conversation

        Raises:
            ConversationError: If the conversation is not found
        """
        if conversation_id not in self.conversations:
            # Try to load from file
            if self.storage_dir:
                try:
                    self._load_conversation(conversation_id)
                except Exception as e:
                    raise ConversationError(f"Conversation not found: {conversation_id}")
            else:
                raise ConversationError(f"Conversation not found: {conversation_id}")

        return self.conversations[conversation_id]

    def add_message(
        self,
        conversation_id: str,
        role: MessageRole,
        content: str
    ) -> Message:
        """
        Add a message to a conversation.

        Args:
            conversation_id: Conversation ID
            role: Message role
            content: Message content

        Returns:
            Added message

        Raises:
            ConversationError: If the conversation is not found
        """
        conversation = self.get_conversation(conversation_id)
        message = conversation.add_message(role, content)
        self._save_conversation(conversation)
        return message

    def list_conversations(self, user_id: str) -> List[Dict]:
        """
        List conversations for a user.

        Args:
            user_id: User ID

        Returns:
            List of conversation summaries
        """
        # Load all conversations for the user
        if self.storage_dir:
            self._load_all_conversations(user_id)

        # Filter conversations by user ID
        user_conversations = [
            conv for conv in self.conversations.values()
            if conv.user_id == user_id
        ]

        # Sort by updated_at (newest first)
        user_conversations.sort(key=lambda c: c.updated_at, reverse=True)

        # Create summaries
        return [
            {
                'id': conv.id,
                'title': conv.title or 'Untitled conversation',
                'created_at': conv.created_at.isoformat(),
                'updated_at': conv.updated_at.isoformat(),
                'message_count': len(conv.messages)
            }
            for conv in user_conversations
        ]

    def delete_conversation(self, conversation_id: str) -> bool:
        """
        Delete a conversation.

        Args:
            conversation_id: Conversation ID

        Returns:
            True if the conversation was deleted, False otherwise
        """
        if conversation_id not in self.conversations:
            return False

        # Remove from memory
        del self.conversations[conversation_id]

        # Remove from storage
        if self.storage_dir:
            file_path = os.path.join(self.storage_dir, f"{conversation_id}.json")
            if os.path.exists(file_path):
                os.remove(file_path)

        logger.info(f"Deleted conversation {conversation_id}")
        return True

    def _save_conversation(self, conversation: Conversation) -> None:
        """
        Save a conversation to file.

        Args:
            conversation: Conversation to save
        """
        if not self.storage_dir:
            return

        file_path = os.path.join(self.storage_dir, f"{conversation.id}.json")

        try:
            # Convert to JSON-serializable dict
            data = conversation.dict()

            # Convert datetime objects to strings
            data['created_at'] = data['created_at'].isoformat()
            data['updated_at'] = data['updated_at'].isoformat()
            for message in data['messages']:
                message['timestamp'] = message['timestamp'].isoformat()

            # Write to file
            with open(file_path, 'w') as f:
                json.dump(data, f)
        except Exception as e:
            logger.error(f"Error saving conversation {conversation.id}: {str(e)}")

    def _load_conversation(self, conversation_id: str) -> None:
        """
        Load a conversation from file.

        Args:
            conversation_id: Conversation ID

        Raises:
            FileNotFoundError: If the conversation file is not found
        """
        if not self.storage_dir:
            raise FileNotFoundError(f"Storage directory not configured")

        file_path = os.path.join(self.storage_dir, f"{conversation_id}.json")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Conversation file not found: {file_path}")

        try:
            # Read from file
            with open(file_path, 'r') as f:
                data = json.load(f)

            # Convert string timestamps to datetime
            data['created_at'] = datetime.fromisoformat(data['created_at'])
            data['updated_at'] = datetime.fromisoformat(data['updated_at'])
            for message in data['messages']:
                message['timestamp'] = datetime.fromisoformat(message['timestamp'])

            # Create conversation
            conversation = Conversation(**data)
            self.conversations[conversation_id] = conversation
        except Exception as e:
            logger.error(f"Error loading conversation {conversation_id}: {str(e)}")
            raise

    def _load_all_conversations(self, user_id: str) -> None:
        """
        Load all conversations for a user.

        Args:
            user_id: User ID
        """
        if not self.storage_dir or not os.path.exists(self.storage_dir):
            return

        for filename in os.listdir(self.storage_dir):
            if filename.endswith('.json'):
                conversation_id = filename[:-5]  # Remove .json extension

                # Skip if already loaded
                if conversation_id in self.conversations:
                    continue

                try:
                    self._load_conversation(conversation_id)

                    # Check if conversation belongs to user
                    conversation = self.conversations.get(conversation_id)
                    if conversation and conversation.user_id != user_id:
                        # Remove from memory if not for this user
                        del self.conversations[conversation_id]
                except Exception as e:
                    logger.error(f"Error loading conversation {conversation_id}: {str(e)}")

# core/user_settings_storage.py
import json
import os
from datetime import datetime
from typing import Dict, Optional

from core.exceptions import SettingsError
from core.logging import get_logger
from core.user_settings import UserSettings

logger = get_logger(__name__)

class UserSettingsStorage:
    """
    Storage for user settings.

    Supports in-memory storage and file-based persistence.
    """

    def __init__(self, storage_dir: Optional[str] = None):
        """
        Initialize the user settings storage.

        Args:
            storage_dir: Directory for file-based storage
        """
        self.settings: Dict[str, UserSettings] = {}
        self.storage_dir = storage_dir

        # Create storage directory if it doesn't exist
        if storage_dir and not os.path.exists(storage_dir):
            os.makedirs(storage_dir)
            logger.info(f"Created user settings storage directory: {storage_dir}")

    def get_settings(self, user_id: str) -> UserSettings:
        """
        Get settings for a user.

        Args:
            user_id: User ID

        Returns:
            User settings
        """
        if user_id not in self.settings:
            # Try to load from file
            if self.storage_dir:
                try:
                    self._load_settings(user_id)
                except Exception as e:
                    # Create new settings
                    self.settings[user_id] = UserSettings(user_id=user_id)
            else:
                # Create new settings
                self.settings[user_id] = UserSettings(user_id=user_id)

        return self.settings[user_id]

    def save_settings(self, settings: UserSettings) -> None:
        """
        Save settings for a user.

        Args:
            settings: User settings
        """
        # Update in memory
        self.settings[settings.user_id] = settings

        # Save to file
        self._save_settings(settings)

    def _save_settings(self, settings: UserSettings) -> None:
        """
        Save settings to file.

        Args:
            settings: User settings
        """
        if not self.storage_dir:
            return

        file_path = os.path.join(self.storage_dir, f"{settings.user_id}.json")

        try:
            # Convert to JSON-serializable dict
            data = settings.dict()

            # Convert datetime objects to strings
            data['created_at'] = data['created_at'].isoformat()
            data['updated_at'] = data['updated_at'].isoformat()

            # Write to file
            with open(file_path, 'w') as f:
                json.dump(data, f)

            logger.debug(f"Saved settings for user {settings.user_id}")
        except Exception as e:
            logger.error(f"Error saving settings for user {settings.user_id}: {str(e)}")

    def _load_settings(self, user_id: str) -> None:
        """
        Load settings from file.

        Args:
            user_id: User ID

        Raises:
            FileNotFoundError: If the settings file is not found
        """
        if not self.storage_dir:
            raise FileNotFoundError(f"Storage directory not configured")

        file_path = os.path.join(self.storage_dir, f"{user_id}.json")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Settings file not found: {file_path}")

        try:
            # Read from file
            with open(file_path, 'r') as f:
                data = json.load(f)

            # Convert string timestamps to datetime
            data['created_at'] = datetime.fromisoformat(data['created_at'])
            data['updated_at'] = datetime.fromisoformat(data['updated_at'])

            # Create settings
            settings = UserSettings(**data)
            self.settings[user_id] = settings

            logger.debug(f"Loaded settings for user {user_id}")
        except Exception as e:
            logger.error(f"Error loading settings for user {user_id}: {str(e)}")
            raise
